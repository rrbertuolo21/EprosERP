#!/usr/bin/env python3
"""Conversor Markdown -> .docx (fiel o suficiente para os manuais de negócio do EprosERP).

Uso:
    python3 md_to_docx.py <arquivo.md | pasta> [<saida_dir>]

Cobre: títulos (#..######), parágrafos, listas (- * / 1.), tabelas markdown (| a | b |),
citações (>), blocos de código (```), negrito **x**, itálico *x*, código inline `x`.
Não depende de pandoc/soffice — só python-docx.
"""
import sys, os, re, glob

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


INLINE_RE = re.compile(r'(\*\*.+?\*\*|`.+?`|\*.+?\*)')


def add_inline(paragraph, text):
    """Adiciona texto com negrito/itálico/código inline a um parágrafo."""
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = paragraph.add_run(part[2:-2]); r.bold = True
        elif part.startswith('`') and part.endswith('`'):
            r = paragraph.add_run(part[1:-1]); r.font.name = 'Courier New'; r.font.size = Pt(9)
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            r = paragraph.add_run(part[1:-1]); r.italic = True
        else:
            paragraph.add_run(part)


def is_table_sep(line):
    return bool(re.match(r'^\s*\|?[\s:|-]+\|?\s*$', line)) and '-' in line


def convert_md(md_path, docx_path):
    with open(md_path, encoding='utf-8') as f:
        lines = f.read().split('\n')

    doc = Document()
    style = doc.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(10.5)

    i = 0
    in_code = False
    code_buf = []
    while i < len(lines):
        line = lines[i]

        # bloco de código
        if line.strip().startswith('```'):
            if in_code:
                p = doc.add_paragraph()
                r = p.add_run('\n'.join(code_buf)); r.font.name = 'Courier New'; r.font.size = Pt(9)
                code_buf = []; in_code = False
            else:
                in_code = True
            i += 1; continue
        if in_code:
            code_buf.append(line); i += 1; continue

        # tabela: linha com | e a próxima é separador
        if line.strip().startswith('|') and i + 1 < len(lines) and is_table_sep(lines[i + 1]):
            header = [c.strip() for c in line.strip().strip('|').split('|')]
            rows = []
            j = i + 2
            while j < len(lines) and lines[j].strip().startswith('|'):
                rows.append([c.strip() for c in lines[j].strip().strip('|').split('|')])
                j += 1
            table = doc.add_table(rows=1, cols=len(header)); table.style = 'Light Grid Accent 1'
            for k, h in enumerate(header):
                cell = table.rows[0].cells[k]; cell.paragraphs[0].clear()
                add_inline(cell.paragraphs[0], h)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            for row in rows:
                cells = table.add_row().cells
                for k in range(len(header)):
                    val = row[k] if k < len(row) else ''
                    cells[k].paragraphs[0].clear(); add_inline(cells[k].paragraphs[0], val)
            doc.add_paragraph()
            i = j; continue

        # títulos
        m = re.match(r'^(#{1,6})\s+(.*)$', line)
        if m:
            level = len(m.group(1)); txt = m.group(2).strip()
            doc.add_heading(txt, level=min(level, 4)); i += 1; continue

        # citação
        if line.strip().startswith('>'):
            p = doc.add_paragraph(); p.style = 'Intense Quote'
            add_inline(p, line.strip().lstrip('>').strip()); i += 1; continue

        # lista
        mlist = re.match(r'^(\s*)([-*]|\d+\.)\s+(.*)$', line)
        if mlist:
            indent = len(mlist.group(1))
            ordered = bool(re.match(r'\d+\.', mlist.group(2)))
            style_name = 'List Number' if ordered else 'List Bullet'
            p = doc.add_paragraph(style=style_name)
            if indent >= 2:
                p.paragraph_format.left_indent = Pt(18 * (indent // 2 + 1))
            add_inline(p, mlist.group(3)); i += 1; continue

        # linha horizontal
        if re.match(r'^\s*---+\s*$', line):
            i += 1; continue

        # parágrafo / vazio
        if line.strip():
            p = doc.add_paragraph(); add_inline(p, line)
        i += 1

    os.makedirs(os.path.dirname(docx_path), exist_ok=True)
    doc.save(docx_path)
    return docx_path


def main():
    if len(sys.argv) < 2:
        print(__doc__); sys.exit(1)
    src = sys.argv[1]
    out_dir = sys.argv[2] if len(sys.argv) > 2 else None

    if os.path.isdir(src):
        mds = glob.glob(os.path.join(src, '**', '*.md'), recursive=True)
    else:
        mds = [src]

    count = 0
    for md in mds:
        base = os.path.splitext(os.path.basename(md))[0]
        if out_dir:
            docx_path = os.path.join(out_dir, base + '.docx')
        else:
            docx_path = os.path.splitext(md)[0] + '.docx'
        try:
            convert_md(md, docx_path)
            count += 1
            print(f'OK  {os.path.basename(docx_path)}')
        except Exception as e:
            print(f'ERRO {md}: {e}')
    print(f'--- {count} arquivo(s) convertido(s) ---')


if __name__ == '__main__':
    main()
